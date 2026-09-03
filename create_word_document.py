r"""
Zero-install Word document creation utility (Python)
Run directly with:
    python create_word_document.py
    OR
    .\bundle\python\python.exe create_word_document.py
r"""
import sys
from pathlib import Path

# Add vendor directory to path if not installed globally
vendor_path = Path(__file__).resolve().parent / "vendor"
if vendor_path.exists():
    sys.path.insert(0, str(vendor_path))

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_report(output_file="report_output.docx", title="Executive Summary", paragraphs=None, table_data=None):
    doc = docx.Document()

    # Title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    intro = doc.add_paragraph()
    run = intro.add_run("Generated automatically using pre-bundled python-docx offline tools.")
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # Paragraphs
    if paragraphs:
        for p_text in paragraphs:
            doc.add_paragraph(p_text)

    # Table
    if table_data:
        headers, rows = table_data
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Shading Accent 1' if 'Light Shading Accent 1' in doc.styles else 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        for idx, h in enumerate(headers):
            hdr_cells[idx].text = h
            hdr_cells[idx].paragraphs[0].runs[0].font.bold = True

        for row in rows:
            row_cells = table.add_row().cells
            for idx, val in enumerate(row):
                row_cells[idx].text = str(val)

    doc.save(output_file)
    print(f"Successfully generated Word document: {output_file}")
    return output_file

if __name__ == "__main__":
    create_word_report(
        output_file="sample_python_report.docx",
        title="AI Automated Project Report",
        paragraphs=[
            "This report was generated without running any pip or npm commands.",
            "All dependencies (python-docx, lxml, openpyxl, xlsxwriter) are fully pre-bundled."
        ],
        table_data=(
            ["Module", "Type", "Status"],
            [
                ["python-docx", "Word Engine", "Ready"],
                ["openpyxl", "Excel Engine", "Ready"],
                ["xlsxwriter", "Excel Formatter", "Ready"],
                ["Cloud SQL Proxy", "Secure Tunnel", "Ready"]
            ]
        )
    )
