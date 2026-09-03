"""
Test script verifying Python Word & Excel libraries:
- python-docx
- openpyxl
- xlsxwriter
Works with bundled Python and vendored fallback.
"""
import sys
from pathlib import Path

# Fallback: support vendor directory if not in global site-packages
base_dir = Path(__file__).resolve().parent
vendor_dir = base_dir / "vendor"
if vendor_dir.exists():
    sys.path.insert(0, str(vendor_dir))

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import openpyxl
import xlsxwriter

print("--- PYTHON OFFICE TOOLS VERIFICATION ---")
print("python-docx loaded:", docx.__version__)
print("openpyxl loaded:", openpyxl.__version__)
print("xlsxwriter loaded:", xlsxwriter.__version__)

# 1. Create a Word Document with python-docx
doc = docx.Document()
title = doc.add_heading("AI Generated Word Document", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph("This document was created automatically by the AI offline bundle with zero installation required.")
p.add_run(" It supports rich formatting, ").bold = True
p.add_run("colored text, ").font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
p.add_run("and tables.")

table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "ID"
hdr_cells[1].text = "Feature"
hdr_cells[2].text = "Status"

data = [
    ("1", "Word Document Creation (.docx)", "Operational"),
    ("2", "Excel Spreadsheet Creation (.xlsx)", "Operational"),
    ("3", "Zero-Install Offline Execution", "Ready"),
]
for id_val, feat, stat in data:
    row_cells = table.add_row().cells
    row_cells[0].text = id_val
    row_cells[1].text = feat
    row_cells[2].text = stat

doc_output = base_dir / "test_output.docx"
doc.save(str(doc_output))
print(f"Generated Word docx: {doc_output} ({doc_output.stat().st_size} bytes)")

# 2. Create an Excel Spreadsheet with xlsxwriter
xls_output = base_dir / "test_output_xlsxwriter.xlsx"
wb = xlsxwriter.Workbook(str(xls_output))
ws = wb.add_worksheet("AI Report")

header_format = wb.add_format({
    'bold': True,
    'font_color': 'white',
    'bg_color': '#1F4E79',
    'border': 1
})
num_format = wb.add_format({'num_format': '$#,##0.00', 'border': 1})
border_format = wb.add_format({'border': 1})

headers = ["ID", "Tool", "Category", "Revenue Impact"]
for col, h in enumerate(headers):
    ws.write(0, col, h, header_format)

items = [
    (1, "python-docx", "Word Processing", 45000),
    (2, "openpyxl", "Excel Read/Write", 62000),
    (3, "xlsxwriter", "Excel Advanced Formatting", 58000),
]
for row_idx, item in enumerate(items, start=1):
    ws.write(row_idx, 0, item[0], border_format)
    ws.write(row_idx, 1, item[1], border_format)
    ws.write(row_idx, 2, item[2], border_format)
    ws.write(row_idx, 3, item[3], num_format)

ws.set_column('A:A', 8)
ws.set_column('B:C', 25)
ws.set_column('D:D', 18)
wb.close()
print(f"Generated Excel with xlsxwriter: {xls_output} ({xls_output.stat().st_size} bytes)")

# 3. Verify openpyxl read back
wb_read = openpyxl.load_workbook(str(xls_output))
print(f"Read back Excel with openpyxl: Sheet '{wb_read.active.title}', Rows: {wb_read.active.max_row}")

# Clean up temporary test files
doc_output.unlink(missing_ok=True)
xls_output.unlink(missing_ok=True)

print("ALL PYTHON OFFICE TOOLS OPERATIONAL AND VERIFIED!")
